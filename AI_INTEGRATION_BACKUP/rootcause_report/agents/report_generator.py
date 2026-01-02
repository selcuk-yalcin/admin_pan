"""
Report Generator Agent - HSG245 Format
Generates professional Word/PDF reports in HSG245 format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict
from pathlib import Path


class ReportGeneratorAgent:
    """
    Generates HSG245-compliant investigation reports
    
    Output formats:
    - Word (.docx) - Editable format
    - PDF (.pdf) - Final distribution format
    """
    
    def __init__(self):
        """Initialize Report Generator"""
        print(f"✅ Report Generator initialized")
    
    def generate_word_report(self, investigation_data: Dict, output_path: str) -> str:
        """
        Generate Word document in HSG245 format
        
        Args:
            investigation_data: Complete investigation results
            output_path: Path to save the report
            
        Returns:
            Path to generated report
        """
        print("\n" + "="*80)
        print("📄 GENERATING HSG245 INVESTIGATION REPORT (WORD)")
        print("="*80)
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Extract data
        part1 = investigation_data.get("part1", {})
        part2 = investigation_data.get("part2", {})
        part3 = investigation_data.get("part3_rca", {})
        
        # === TITLE PAGE ===
        self._add_title_page(doc, part1)
        doc.add_page_break()
        
        # === PART 1: OVERVIEW ===
        self._add_part1_overview(doc, part1)
        doc.add_page_break()
        
        # === PART 2: INITIAL ASSESSMENT ===
        self._add_part2_assessment(doc, part2)
        doc.add_page_break()
        
        # === PART 3: ROOT CAUSE ANALYSIS ===
        self._add_part3_root_cause(doc, part3)
        doc.add_page_break()
        
        # === PART 4: RECOMMENDATIONS ===
        self._add_part4_recommendations(doc, part3)
        
        # Save document
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        print(f"✅ Word report generated: {output_path}")
        return output_path
    
    def _add_title_page(self, doc: Document, part1: Dict):
        """Add title page"""
        # Title
        title = doc.add_heading("HSG245 INVESTIGATION REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Subtitle
        subtitle = doc.add_heading("Adverse Event Report and Investigation", 2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Reference table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = "Reference Number:"
        cells[1].text = part1.get("ref_no", "N/A")
        
        cells = table.rows[1].cells
        cells[0].text = "Incident Date/Time:"
        cells[1].text = part1.get("date_time", "N/A")
        
        cells = table.rows[2].cells
        cells[0].text = "Incident Type:"
        cells[1].text = part1.get("incident_type", "N/A")
        
        cells = table.rows[3].cells
        cells[0].text = "Reported By:"
        cells[1].text = part1.get("reported_by", "N/A")
        
        cells = table.rows[4].cells
        cells[0].text = "Report Date:"
        cells[1].text = datetime.now().strftime("%d/%m/%Y")
        
        cells = table.rows[5].cells
        cells[0].text = "Status:"
        cells[1].text = "Investigation Complete"
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Footer
        footer = doc.add_paragraph("Health and Safety Executive")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.runs[0]
        footer_run.italic = True
    
    def _add_part1_overview(self, doc: Document, part1: Dict):
        """Add Part 1: Overview section"""
        # Heading
        doc.add_heading("PART 1: OVERVIEW", 1)
        
        doc.add_paragraph(
            "The purpose of this form is to record all adverse events. "
            "Part 1 should be filled out immediately by the manager or supervisor."
        ).italic = True
        
        doc.add_paragraph()
        
        # Basic information table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = "Reported by:"
        cells[1].text = part1.get("reported_by", "")
        
        cells = table.rows[1].cells
        cells[0].text = "Date/time of adverse event:"
        cells[1].text = part1.get("date_time", "")
        
        cells = table.rows[2].cells
        cells[0].text = "Incident type:"
        cells[1].text = part1.get("incident_type", "")
        
        doc.add_paragraph()
        
        # Brief details
        doc.add_heading("Brief Details", 2)
        
        brief = part1.get("brief_details", {})
        
        p = doc.add_paragraph()
        p.add_run("What: ").bold = True
        p.add_run(brief.get("what", "N/A"))
        
        p = doc.add_paragraph()
        p.add_run("Where: ").bold = True
        p.add_run(brief.get("where", "N/A"))
        
        p = doc.add_paragraph()
        p.add_run("When: ").bold = True
        p.add_run(brief.get("when", "N/A"))
        
        p = doc.add_paragraph()
        p.add_run("Who: ").bold = True
        who_value = brief.get("who", "N/A")
        # Handle dict or list format
        if isinstance(who_value, (dict, list)):
            p.add_run(str(who_value))
        else:
            p.add_run(who_value)
        
        p = doc.add_paragraph()
        p.add_run("Emergency measures taken: ").bold = True
        p.add_run(brief.get("emergency_measures", "N/A"))
        
        doc.add_paragraph()
        
        # Forwarding information
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = "Forwarded to:"
        cells[1].text = part1.get("forwarded_to", "")
        
        cells = table.rows[1].cells
        cells[0].text = "Date/Time:"
        cells[1].text = part1.get("forwarded_date_time", "")
    
    def _add_part2_assessment(self, doc: Document, part2: Dict):
        """Add Part 2: Initial Assessment section"""
        doc.add_heading("PART 2: INITIAL ASSESSMENT", 1)
        
        doc.add_paragraph(
            "To be carried out by the person responsible for health and safety"
        ).italic = True
        
        doc.add_paragraph()
        
        # Event classification table
        doc.add_heading("Event Classification", 2)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = "Type of event:"
        cells[1].text = part2.get("type_of_event", "")
        
        cells = table.rows[1].cells
        cells[0].text = "Actual/potential for harm:"
        cells[1].text = part2.get("actual_potential_harm", "")
        
        cells = table.rows[2].cells
        cells[0].text = "RIDDOR reportable:"
        riddor = part2.get("riddor_reportable", "N")
        cells[1].text = f"{riddor} {'(' + part2.get('riddor_date_reported', '') + ')' if riddor == 'Y' else ''}"
        
        cells = table.rows[3].cells
        cells[0].text = "Accident book entry:"
        book = part2.get("accident_book_entry", "N")
        cells[1].text = f"{book} {part2.get('accident_book_ref', '') if book == 'Y' else ''}"
        
        doc.add_paragraph()
        
        # Investigation details
        doc.add_heading("Investigation Requirements", 2)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = "Investigation level:"
        cells[1].text = part2.get("investigation_level", "")
        
        cells = table.rows[1].cells
        cells[0].text = "Priority:"
        cells[1].text = part2.get("priority", "")
        
        cells = table.rows[2].cells
        cells[0].text = "Further investigation required:"
        cells[1].text = part2.get("further_investigation_required", "")
        
        cells = table.rows[3].cells
        cells[0].text = "Investigation team:"
        team = part2.get("investigation_team", [])
        cells[1].text = ", ".join(team) if team else "N/A"
        
        cells = table.rows[4].cells
        cells[0].text = "Assessed by:"
        cells[1].text = f"{part2.get('initial_assessment_by', '')} on {part2.get('assessment_date', '')}"
    
    def _add_part3_root_cause(self, doc: Document, part3: Dict):
        """Add Part 3: Root Cause Analysis section"""
        doc.add_heading("PART 3: ROOT CAUSE ANALYSIS", 1)
        
        doc.add_heading("Analysis Method: 5 Why Analysis", 2)
        
        # 5 Why Chains
        chains = part3.get("five_why_chains", [])
        
        for chain in chains:
            chain_id = chain.get("chain_id", "?")
            description = chain.get("description", "N/A")
            
            doc.add_heading(f"Causal Chain {chain_id}", 3)
            doc.add_paragraph(description).italic = True
            
            whys = chain.get("whys", [])
            for idx, why in enumerate(whys, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"Why {idx}: ").bold = True
                p.add_run(why.get("question", "N/A"))
                
                p = doc.add_paragraph(style='List Continue')
                p.add_run("→ ").bold = True
                p.add_run(why.get("answer", "N/A"))
            
            # Root cause box
            root_cause = chain.get("root_cause", "N/A")
            p = doc.add_paragraph()
            p.add_run("🎯 ROOT CAUSE: ").bold = True
            run = p.add_run(root_cause)
            run.font.color.rgb = RGBColor(255, 0, 0)
            
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # Categorized causes
        doc.add_heading("Categorized Causes", 2)
        
        # Immediate causes
        doc.add_heading("⚡ Immediate Causes", 3)
        immediate = part3.get("immediate_causes", [])
        for cause in immediate:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(cause.get("cause", "N/A")).bold = True
            if cause.get("description"):
                p = doc.add_paragraph(style='List Continue')
                p.add_run(cause["description"])
        
        doc.add_paragraph()
        
        # Underlying causes
        doc.add_heading("🔧 Underlying Causes", 3)
        underlying = part3.get("underlying_causes", [])
        for cause in underlying:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(cause.get("cause", "N/A")).bold = True
            if cause.get("description"):
                p = doc.add_paragraph(style='List Continue')
                p.add_run(cause["description"])
        
        doc.add_paragraph()
        
        # Root causes
        doc.add_heading("🎯 Root Causes (Systemic)", 3)
        root = part3.get("root_causes", [])
        for cause in root:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(cause.get("cause", "N/A"))
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            if cause.get("description"):
                p = doc.add_paragraph(style='List Continue')
                p.add_run(cause["description"])
    
    def _add_part4_recommendations(self, doc: Document, part3: Dict):
        """Add Part 4: Recommendations section"""
        doc.add_heading("PART 4: RECOMMENDATIONS & ACTION PLAN", 1)
        
        doc.add_paragraph(
            "Risk control measures needed based on root cause analysis findings"
        ).italic = True
        
        doc.add_paragraph()
        
        # Generate recommendations based on root causes
        doc.add_heading("Recommended Control Measures", 2)
        
        root_causes = part3.get("root_causes", [])
        
        table = doc.add_table(rows=len(root_causes) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Root Cause"
        header_cells[1].text = "Control Measure"
        header_cells[2].text = "Priority"
        header_cells[3].text = "Responsible"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Recommendations
        for idx, cause in enumerate(root_causes, 1):
            cells = table.rows[idx].cells
            cells[0].text = cause.get("cause", "N/A")[:50]
            cells[1].text = f"Address {cause.get('cause', 'issue')}"
            cells[2].text = "High" if idx <= 2 else "Medium"
            cells[3].text = "H&S Officer"
        
        doc.add_paragraph()
        
        # Investigation team sign-off
        doc.add_heading("Investigation Team", 2)
        
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = "Name"
        header_cells[1].text = "Position"
        header_cells[2].text = "Signature"
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Team members (placeholder)
        team_members = [
            ("H&S Officer", "Health & Safety Officer"),
            ("Line Manager", "Operations Manager"),
            ("Technical Expert", "Engineering Specialist")
        ]
        
        for idx, (name, position) in enumerate(team_members, 1):
            cells = table.rows[idx].cells
            cells[0].text = name
            cells[1].text = position
            cells[2].text = "_________________"
        
        doc.add_paragraph()
        
        # Date
        p = doc.add_paragraph()
        p.add_run("Report Date: ").bold = True
        p.add_run(datetime.now().strftime("%d/%m/%Y %H:%M"))
